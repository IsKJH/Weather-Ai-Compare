from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

OUTPUT = r"C:\proj\weather-ai-compare\results\AI_Weather_App_비교_리포트.docx"
SCREENSHOTS = r"C:\proj\weather-ai-compare\screenshots"

# ── 색상 팔레트 ──────────────────────────────────────────────
C_BLUE   = RGBColor(0x1A, 0x73, 0xE8)
C_DARK   = RGBColor(0x20, 0x20, 0x2E)
C_GRAY   = RGBColor(0x5F, 0x6B, 0x7C)
C_LIGHT  = RGBColor(0xF8, 0xF9, 0xFA)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_GREEN  = RGBColor(0x1E, 0x8E, 0x3E)
C_ORANGE = RGBColor(0xFA, 0x7B, 0x17)
C_RED    = RGBColor(0xD9, 0x30, 0x25)

CLAUDE_COLOR = RGBColor(0xD9, 0x7B, 0x2B)
CODEX_COLOR  = RGBColor(0x10, 0xA3, 0x7F)
GEMINI_COLOR = RGBColor(0x42, 0x85, 0xF4)


# ── 헬퍼 함수 ────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{ side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_paragraph(doc, text='', bold=False, size=11, color=None, align=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align: p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color: run.font.color.rgb = color
    return p

def add_heading(doc, text, level=1, color=C_DARK):
    sizes = {1: 20, 2: 15, 3: 13}
    p = add_paragraph(doc, text, bold=True, size=sizes.get(level, 12),
                      color=color, space_before=14, space_after=6)
    if level == 1:
        # 하단 선 추가
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), '1A73E8')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def colored_run(para, text, bold=False, size=11, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return run


# ── 문서 생성 ────────────────────────────────────────────────
doc = Document()

# 여백 설정
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# 기본 폰트
doc.styles['Normal'].font.name = '맑은 고딕'
doc.styles['Normal'].font.size = Pt(10.5)


# ════════════════════════════════════════════════════════════
# 1. 표지
# ════════════════════════════════════════════════════════════
add_paragraph(doc, space_before=40, space_after=0)

p = add_paragraph(doc, 'AI 코딩 도구 비교 실험 리포트', bold=True, size=28,
                  color=C_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=8)

add_paragraph(doc, 'Claude  vs  Codex  vs  Gemini', bold=True, size=16,
              color=C_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)

add_paragraph(doc, '동일한 Android 날씨 앱을 3개의 AI로 만들면 무엇이 달라지는가?', size=12,
              color=C_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=40)

add_paragraph(doc, '2026년 5월', size=11, color=C_GRAY,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 2. 목차
# ════════════════════════════════════════════════════════════
add_heading(doc, '목  차')
toc_items = [
    ('1', '프로젝트 기획 의도'),
    ('2', '실험 설계'),
    ('3', '수치 비교'),
    ('4', '토큰 사용량 상세 분석'),
    ('5', 'UI 스크린샷 비교'),
    ('6', '코드 스타일 분석'),
    ('7', 'AI별 장점과 단점'),
    ('8', '결론 및 용도별 추천'),
]
for num, title in toc_items:
    p = add_paragraph(doc, space_before=2, space_after=2)
    colored_run(p, f'  {num}.  ', bold=True, color=C_BLUE)
    colored_run(p, title, color=C_DARK)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 3. 기획 의도
# ════════════════════════════════════════════════════════════
add_heading(doc, '1.  프로젝트 기획 의도')

add_paragraph(doc,
    'AI 코딩 도구가 빠르게 확산되면서 개발자들은 Claude, ChatGPT(Codex), Gemini 중 어떤 도구를 '
    '선택해야 하는지 막막한 상황입니다. 이 실험은 동일한 조건에서 세 AI에게 같은 과제를 주고 '
    '실제로 어떤 차이가 나타나는지를 객관적인 수치와 결과물로 비교합니다.',
    size=11, color=C_DARK, space_after=8)

# 핵심 질문 박스
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = tbl.cell(0, 0)
set_cell_bg(cell, 'EBF3FB')
cell.paragraphs[0].clear()
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(8)
p.paragraph_format.left_indent  = Cm(0.5)
colored_run(p, '핵심 질문  ', bold=True, size=11, color=C_BLUE)
colored_run(p, '같은 앱을 만들 때 토큰 비용, 속도, 코드 품질, UI 디자인이 AI마다 얼마나 다른가?',
            size=11, color=C_DARK)
add_paragraph(doc, space_before=8, space_after=0)

add_heading(doc, '비교 관점', level=2)
points = [
    ('토큰 사용량', '각 API 호출 비용의 직접적 지표'),
    ('생성 속도',   '같은 작업에 걸리는 시간'),
    ('코드 구조',   '파일 수, 라인 수, 아키텍처 선택'),
    ('UI 디자인',   '같은 명세에서 어떤 시각적 차이가 나타나는가'),
    ('자동화 수준', '빌드, 오류 처리를 스스로 해결하는 정도'),
    ('한국어 처리', '한국어 데이터 처리 능력'),
]
for title, desc in points:
    p = add_paragraph(doc, space_before=2, space_after=2)
    colored_run(p, f'  ▪  {title}  —  ', bold=True, size=10.5, color=C_BLUE)
    colored_run(p, desc, size=10.5, color=C_DARK)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 4. 실험 설계
# ════════════════════════════════════════════════════════════
add_heading(doc, '2.  실험 설계')

add_heading(doc, '고정 조건 (3개 AI 동일)', level=2)

tbl = doc.add_table(rows=7, cols=2)
tbl.style = 'Table Grid'
headers = [('항목', 'E8F0FE'), ('값', 'E8F0FE')]
rows_data = [
    ('과제',        'Android 날씨 앱 (WeatherNow) 제작'),
    ('언어',        'Kotlin + Jetpack Compose'),
    ('아키텍처',    'MVVM (ViewModel + StateFlow)'),
    ('데이터',      'Mock 데이터 (서울/부산/제주 — 완전히 동일)'),
    ('외부 라이브러리', '사용 금지'),
    ('프롬프트',    '3개 AI 모두 완전히 동일한 텍스트'),
]
for i, (h, bg) in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    set_cell_bg(cell, bg.replace('#',''))
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = C_DARK

for r_idx, (k, v) in enumerate(rows_data, start=1):
    tbl.rows[r_idx].cells[0].text = k
    tbl.rows[r_idx].cells[1].text = v
    for c in tbl.rows[r_idx].cells:
        for para in c.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10.5)

tbl.columns[0].width = Cm(4)
tbl.columns[1].width = Cm(11)
add_paragraph(doc, space_before=8, space_after=0)

add_heading(doc, '자유 조건 (AI 개성 발휘 영역)', level=2)
free = ['색상 테마 및 배경', '레이아웃 구성', '날씨 아이콘 스타일', '카드/섹션 디자인', '애니메이션 유무']
for item in free:
    p = add_paragraph(doc, f'  ✔  {item}', size=10.5, color=C_DARK, space_before=1, space_after=1)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 5. 수치 비교
# ════════════════════════════════════════════════════════════
add_heading(doc, '3.  수치 비교')

tbl = doc.add_table(rows=8, cols=4)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

header_data = ['항목', 'Claude', 'Codex', 'Gemini']
header_bg   = ['2D2D3E', 'D97B2B', '10A37F', '4285F4']

for c_idx, (text, bg) in enumerate(zip(header_data, header_bg)):
    cell = tbl.rows[0].cells[c_idx]
    set_cell_bg(cell, bg)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = C_WHITE

data_rows = [
    ('사용 모델',       'claude-sonnet-4-6',  'gpt-5.5',       'gemini-3-flash-preview'),
    ('생성 소요 시간',  '537초 (9분)',         '측정 중',        '318초 (5.3분)'),
    ('KT 파일 수',      '5개',                 '6개',            '5개'),
    ('KT 코드 라인 수', '395줄',               '451줄',          '304줄'),
    ('빌드 성공',       '✔ 성공',              '✔ 성공',         '✔ 성공'),
    ('APK 자동 빌드',   '✘',                   '✔ (자동)',       '✘'),
    ('한국어 처리',     '✔ 정상',              '✘ 깨짐',         '✔ 정상'),
]
alt_bg = 'F8F9FA'
for r_idx, row in enumerate(data_rows, start=1):
    bg = alt_bg if r_idx % 2 == 0 else 'FFFFFF'
    for c_idx, text in enumerate(row):
        cell = tbl.rows[r_idx].cells[c_idx]
        if c_idx > 0: set_cell_bg(cell, bg)
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        if c_idx == 0: run.bold = True
        # 색상 강조
        if '✔' in text: run.font.color.rgb = C_GREEN
        elif '✘' in text: run.font.color.rgb = C_RED

for i, w in enumerate([Cm(4.5), Cm(3.5), Cm(3.5), Cm(4)]):
    tbl.columns[i].width = w

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 6. 토큰 상세
# ════════════════════════════════════════════════════════════
add_heading(doc, '4.  토큰 사용량 상세 분석')

# Claude
add_heading(doc, 'Claude  (claude-sonnet-4-6)', level=2, color=CLAUDE_COLOR)
tbl = doc.add_table(rows=6, cols=2)
tbl.style = 'Table Grid'
claude_rows = [
    ('신규 입력 토큰',      '40'),
    ('캐시 생성 토큰',      '12,305'),
    ('캐시 읽기 토큰',      '865,307  ← 프롬프트 캐싱'),
    ('출력 토큰',           '40,626'),
    ('총 입력 합계',        '877,652'),
    ('총합 (입력 + 출력)',  '918,278'),
]
for r_idx, (k, v) in enumerate(claude_rows):
    tbl.rows[r_idx].cells[0].text = k
    tbl.rows[r_idx].cells[1].text = v
    for c_idx, cell in enumerate(tbl.rows[r_idx].cells):
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10.5)
                if r_idx == 5: run.bold = True
        if r_idx == 5: set_cell_bg(cell, 'FFF3CD')
tbl.columns[0].width = Cm(7); tbl.columns[1].width = Cm(8.5)

p = add_paragraph(doc, space_before=4, space_after=4)
colored_run(p, '  ※ ', bold=True, color=C_ORANGE)
colored_run(p, '캐시 읽기 토큰은 일반 입력 대비 약 10배 저렴하게 과금. 실제 비용은 숫자보다 훨씬 낮음.',
            size=10, color=C_GRAY)

# Codex
add_heading(doc, 'Codex  (gpt-5.5)', level=2, color=CODEX_COLOR)
tbl = doc.add_table(rows=2, cols=2)
tbl.style = 'Table Grid'
codex_rows = [
    ('총 사용 토큰 (합계)',    '165,761'),
    ('입력 / 출력 분리',       'CLI에서 미제공 — 합계만 출력됨'),
]
for r_idx, (k, v) in enumerate(codex_rows):
    tbl.rows[r_idx].cells[0].text = k
    tbl.rows[r_idx].cells[1].text = v
    for cell in tbl.rows[r_idx].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10.5)
                if r_idx == 0: run.bold = True
    if r_idx == 0: set_cell_bg(tbl.rows[r_idx].cells[1], 'D1FAE5')
tbl.columns[0].width = Cm(7); tbl.columns[1].width = Cm(8.5)

p = add_paragraph(doc, space_before=4, space_after=4)
colored_run(p, '  ※ ', bold=True, color=C_ORANGE)
colored_run(p, 'Codex CLI는 세션 종료 시 총 토큰만 표시. 3개 중 가장 적은 토큰으로 완성.',
            size=10, color=C_GRAY)

# Gemini
add_heading(doc, 'Gemini  (gemini-3-flash-preview)', level=2, color=GEMINI_COLOR)
tbl = doc.add_table(rows=7, cols=2)
tbl.style = 'Table Grid'
gemini_rows = [
    ('메인 모델',                      'gemini-3-flash-preview'),
    ('입력 토큰',                      '106,847'),
    ('출력 토큰',                      '3,521'),
    ('캐시 토큰',                      '555,670'),
    ('추론(Thinking) 토큰',            '4,122'),
    ('메인 모델 총계',                 '670,160'),
    ('라우터 모델 (flash-lite) 추가',  '+2,362  →  총합 672,522'),
]
for r_idx, (k, v) in enumerate(gemini_rows):
    tbl.rows[r_idx].cells[0].text = k
    tbl.rows[r_idx].cells[1].text = v
    for cell in tbl.rows[r_idx].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10.5)
                if r_idx >= 5: run.bold = True
    if r_idx >= 5: set_cell_bg(tbl.rows[r_idx].cells[1], 'E8F0FE')
tbl.columns[0].width = Cm(7); tbl.columns[1].width = Cm(8.5)

p = add_paragraph(doc, space_before=4, space_after=4)
colored_run(p, '  ※ ', bold=True, color=C_ORANGE)
colored_run(p, 'Gemini는 라우터 + 메인 2개 모델을 동시 사용. Thinking 토큰은 내부 추론에 사용된 토큰.',
            size=10, color=C_GRAY)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 7. 스크린샷 비교
# ════════════════════════════════════════════════════════════
add_heading(doc, '5.  UI 스크린샷 비교')

add_paragraph(doc, '동일한 Mock 데이터(서울, 23°C, 맑음)를 표시하는 화면. 프롬프트는 동일하지만 UI 디자인은 완전히 다르게 나타남.',
              size=10.5, color=C_GRAY, space_after=10)

tbl = doc.add_table(rows=2, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
ai_labels = [('Claude', CLAUDE_COLOR), ('Codex', CODEX_COLOR), ('Gemini', GEMINI_COLOR)]
screenshots = ['claude.png', 'codex.png', 'gemini.png']

for c_idx, (label, color) in enumerate(ai_labels):
    cell = tbl.rows[0].cells[c_idx]
    set_cell_bg(cell, 'FFFFFF')
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.bold = True; run.font.size = Pt(12)
    run.font.color.rgb = color

for c_idx, fname in enumerate(screenshots):
    cell = tbl.rows[1].cells[c_idx]
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = os.path.join(SCREENSHOTS, fname)
    if os.path.exists(img_path):
        run = p.add_run()
        run.add_picture(img_path, width=Cm(4.5))
    else:
        p.add_run(f'[{fname} 없음]')

for col in tbl.columns:
    col.width = Cm(5.2)

add_paragraph(doc, space_before=10, space_after=4)

# UI 특징 비교
add_heading(doc, 'UI 특징 비교', level=2)
tbl2 = doc.add_table(rows=6, cols=4)
tbl2.style = 'Table Grid'
ui_headers = ['UI 요소', 'Claude', 'Codex', 'Gemini']
ui_header_bg = ['2D2D3E', 'D97B2B', '10A37F', '4285F4']
for c_idx, (text, bg) in enumerate(zip(ui_headers, ui_header_bg)):
    cell = tbl2.rows[0].cells[c_idx]
    set_cell_bg(cell, bg)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text); run.bold = True
    run.font.size = Pt(10.5); run.font.color.rgb = C_WHITE

ui_data = [
    ('배경 테마',    '하늘색 그라데이션',  '다크 네이비',       '인디고/보라 그라데이션'),
    ('도시 선택 UI', '상단 탭 (pill)',      '상단 탭 버튼',      '상단 탭'),
    ('날씨 아이콘',  '이모지 (대형)',       '이모지 (깨짐)',      '이모지'),
    ('상세 정보',    '카드 3개',            '2열 배치',          '카드형'),
    ('5일 예보',     '세로 리스트',         '세로 리스트',       '세로 리스트'),
]
for r_idx, row in enumerate(ui_data, start=1):
    bg = 'F8F9FA' if r_idx % 2 == 0 else 'FFFFFF'
    for c_idx, text in enumerate(row):
        cell = tbl2.rows[r_idx].cells[c_idx]
        if c_idx > 0: set_cell_bg(cell, bg)
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        run = p.add_run(text); run.font.size = Pt(10.5)
        if c_idx == 0: run.bold = True

for i, w in enumerate([Cm(3.5), Cm(4), Cm(4), Cm(4)]):
    tbl2.columns[i].width = w

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 8. 코드 분석
# ════════════════════════════════════════════════════════════
add_heading(doc, '6.  코드 스타일 분석')

analyses = [
    ('Claude', CLAUDE_COLOR, [
        ('파일 구조', 'data/, ui/, ui/theme/, viewmodel/ 로 역할별 패키지 분리'),
        ('특징',      'libs.versions.toml 버전 카탈로그 사용 (최신 Gradle 방식)'),
        ('UI 코드',   'WeatherScreen.kt 단일 파일에 전체 UI — 컴포넌트 분리 없음'),
        ('추가 사항', 'edge-to-edge 지원, enableEdgeToEdge() 적용'),
        ('단점',      'gradlew, gradle-wrapper.jar 생성 누락'),
    ]),
    ('Codex', CODEX_COLOR, [
        ('파일 구조', 'ui/, ui/components/ 로 컴포넌트 분리 (ForecastRow, MetricTile 별도)'),
        ('특징',      '생성 완료 후 자동으로 gradlew assembleDebug 실행 → APK 생성'),
        ('UI 코드',   '재사용 컴포저블을 components/ 패키지로 분리하여 구조화'),
        ('추가 사항', 'lifecycle-viewmodel-ktx 외부 라이브러리 추가 (명세 위반)'),
        ('단점',      '한국어 인코딩 깨짐 — Mock 데이터의 한글이 ?? 로 표시됨'),
    ]),
    ('Gemini', GEMINI_COLOR, [
        ('파일 구조', 'ui/, ui/components/ 로 WeatherComponents.kt 단일 컴포넌트 파일'),
        ('특징',      '가장 적은 코드(304줄)로 요구사항 구현 — 간결한 스타일'),
        ('UI 코드',   'WeatherScreen + WeatherComponents 2파일 분리'),
        ('추가 사항', '빌드 자체 검증 실행 (Build Status: Verified 명시)'),
        ('단점',      'mipmap 리소스 누락 — 수동 보완 없이는 빌드 실패'),
    ]),
]
for ai, color, items in analyses:
    add_heading(doc, ai, level=2, color=color)
    tbl = doc.add_table(rows=len(items), cols=2)
    tbl.style = 'Table Grid'
    for r_idx, (k, v) in enumerate(items):
        tbl.rows[r_idx].cells[0].text = k
        tbl.rows[r_idx].cells[1].text = v
        for cell in tbl.rows[r_idx].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size = Pt(10.5)
        tbl.rows[r_idx].cells[0].paragraphs[0].runs[0].bold = True
    tbl.columns[0].width = Cm(3); tbl.columns[1].width = Cm(12.5)
    add_paragraph(doc, space_before=6, space_after=0)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 9. 장단점
# ════════════════════════════════════════════════════════════
add_heading(doc, '7.  AI별 장점과 단점')

pros_cons = [
    ('Claude', CLAUDE_COLOR, 'D97B2B',
     ['최신 Gradle 방식(libs.versions.toml) 사용',
      '날씨별 동적 배경색 등 세련된 UI',
      '한국어 처리 정상',
      '프롬프트 캐싱으로 실질 비용 절감'],
     ['gradlew/jar 파일 미생성 (수동 보완 필요)',
      '비대화형 모드에서 권한 요청 발생',
      '단일 파일 UI (컴포넌트 분리 미흡)']),
    ('Codex', CODEX_COLOR, '10A37F',
     ['생성 후 자동 빌드 + APK 생성',
      '컴포넌트를 별도 파일로 분리',
      '가장 많은 코드(451줄)로 상세 구현',
      '3개 중 실질 사용 토큰 최소(165,761)'],
     ['한국어 인코딩 깨짐 (stdin 파이프 한계)',
      '명세 위반 — 외부 라이브러리 추가',
      '입력/출력 토큰 분리 데이터 미제공']),
    ('Gemini', GEMINI_COLOR, '4285F4',
     ['가장 빠른 생성 (318초)',
      '간결한 코드(304줄)로 핵심 구현',
      '2개 모델 협력 (라우터 + 메인)',
      '한국어 처리 정상'],
     ['mipmap 리소스 누락 (빌드 실패)',
      '비대화형 파이프 모드에서 PTY 오류',
      'Thinking 토큰 별도 발생 (비용 증가)']),
]

for ai, color, bg_hex, pros, cons in pros_cons:
    add_heading(doc, ai, level=2, color=color)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'

    for c_idx, (label, items, bg) in enumerate([
        ('장점 ✔', pros, 'E6F4EA'),
        ('단점 ✘', cons, 'FCE8E6'),
    ]):
        cell = tbl.rows[0].cells[c_idx]
        set_cell_bg(cell, bg)
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.bold = True; run.font.size = Pt(11)
        run.font.color.rgb = C_GREEN if '✔' in label else C_RED
        for item in items:
            p2 = cell.add_paragraph(f'   • {item}')
            for run2 in p2.runs: run2.font.size = Pt(10.5)

    tbl.columns[0].width = Cm(7.5); tbl.columns[1].width = Cm(8)
    add_paragraph(doc, space_before=6, space_after=0)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 10. 결론
# ════════════════════════════════════════════════════════════
add_heading(doc, '8.  결론 및 용도별 추천')

add_paragraph(doc,
    '3개의 AI 모두 같은 명세서로 동작하는 Android 앱을 완성했습니다. '
    '하지만 접근 방식, 속도, 비용, 코드 품질에서 뚜렷한 차이를 보였습니다.',
    size=11, color=C_DARK, space_after=10)

rec_tbl = doc.add_table(rows=4, cols=3)
rec_tbl.style = 'Table Grid'
rec_headers = ['추천 상황', 'AI', '이유']
rec_bg = ['2D2D3E', '2D2D3E', '2D2D3E']
for c_idx, (text, bg) in enumerate(zip(rec_headers, rec_bg)):
    cell = rec_tbl.rows[0].cells[c_idx]
    set_cell_bg(cell, bg)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text); run.bold = True
    run.font.size = Pt(10.5); run.font.color.rgb = C_WHITE

rec_data = [
    ('코드 완성도 + 최신 패턴',  'Claude',  '최신 Gradle 방식, 세련된 UI, 캐싱으로 비용 효율'),
    ('빠른 프로토타입 + 자동화', 'Codex',   '자동 빌드, APK 생성까지 원스톱, 최소 토큰'),
    ('속도 우선 + 간결한 코드',  'Gemini',  '가장 빠른 생성, 깔끔한 코드, Thinking 기반 추론'),
]
rec_bgs = ['FFF3CD', 'D1FAE5', 'DBEAFE']
for r_idx, (situation, ai, reason) in enumerate(rec_data, start=1):
    for c_idx, text in enumerate([situation, ai, reason]):
        cell = rec_tbl.rows[r_idx].cells[c_idx]
        if c_idx == 1: set_cell_bg(cell, rec_bgs[r_idx - 1])
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(10.5)
        if c_idx == 1: run.bold = True

for i, w in enumerate([Cm(5), Cm(2.5), Cm(8)]):
    rec_tbl.columns[i].width = w

add_paragraph(doc, space_before=14, space_after=4)
add_heading(doc, '핵심 인사이트', level=2)

insights = [
    '토큰 수치만으로 비용을 비교하면 안 됩니다 — Claude의 캐시 읽기, Gemini의 Thinking 토큰은 과금 방식이 다릅니다.',
    'Codex는 가장 적은 토큰(165,761)으로 APK까지 완성했지만, 한국어 처리 한계가 실무 적용의 걸림돌입니다.',
    '3개 AI 모두 같은 요구사항을 충족했지만 UI 디자인은 완전히 달랐습니다 — AI에게도 "디자인 개성"이 있습니다.',
    'Gemini는 2개 모델을 협력시키는 구조로, 단순 토큰 비교 이상의 복잡성이 있습니다.',
]
for insight in insights:
    p = add_paragraph(doc, space_before=3, space_after=3)
    colored_run(p, '  ►  ', bold=True, color=C_BLUE)
    colored_run(p, insight, size=10.5, color=C_DARK)

# 푸터
add_paragraph(doc, space_before=30, space_after=0)
add_paragraph(doc, 'AI 코딩 도구 비교 실험 리포트  |  2026년 5월',
              size=9, color=C_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)


# ── 저장 ────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"저장 완료: {OUTPUT}")
